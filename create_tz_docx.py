from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Voice_AI_Companion_TZ.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(96, 106, 116)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = RGBColor(31, 99, 61)
RED = RGBColor(155, 28, 28)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Pt(widths[i] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def paragraph_border_bottom(paragraph, color="2E74B5", size="8", space="6"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_para(doc, text="", style=None, size=11, color=None, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = add_para(doc, text, style=style, after=4)
    return p


def add_number(doc, text):
    return add_para(doc, text, style="List Number", after=4)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt({1: 18, 2: 14, 3: 10}.get(level, 8))
    p.paragraph_format.space_after = Pt({1: 10, 2: 7, 3: 5}.get(level, 4))
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for r in p.runs:
        set_run_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11),
                     color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_callout(doc, title, body, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, color=INK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=RGBColor(40, 45, 50))
    add_para(doc, "", after=4)


def add_kv_table(doc, rows, widths=(2200, 7160), header=None):
    table = doc.add_table(rows=0, cols=2)
    set_table_width(table, list(widths))
    for idx, (k, v) in enumerate(rows):
        cells = table.add_row().cells
        if header and idx == 0:
            pass
        set_cell_shading(cells[0], LIGHT_BLUE)
        for p in cells[0].paragraphs:
            p.paragraph_format.space_after = Pt(0)
        r = cells[0].paragraphs[0].add_run(k)
        set_run_font(r, size=10, color=INK, bold=True)
        for p in cells[1].paragraphs:
            p.paragraph_format.space_after = Pt(0)
        r2 = cells[1].paragraphs[0].add_run(v)
        set_run_font(r2, size=10, color=RGBColor(30, 30, 30))
    add_para(doc, "", after=4)
    return table


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(value))
            set_run_font(r, size=9.3, color=RGBColor(28, 28, 28))
    add_para(doc, "", after=6)
    return table


def set_document_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25

    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header_p.add_run("Voice-Only AI Companion TZ")
    set_run_font(r, size=9, color=MUTED, bold=True)
    paragraph_border_bottom(header_p, color="DADCE0", size="4", space="4")

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer_p.add_run("Version 1.0  |  macOS MVP")
    set_run_font(r, size=9, color=MUTED)


def add_cover(doc):
    add_para(doc, "TECHNICAL SPECIFICATION", size=10.5, color=MUTED, bold=True, after=8)
    title = add_para(doc, "Voice-Only AI Companion", size=25, color=RGBColor(0, 0, 0), bold=True, after=4)
    title.paragraph_format.line_spacing = 1.0
    add_para(doc, "macOS uchun Unreal/MetaHuman darajasidagi ovozli AI avatar platformasi", size=14, color=RGBColor(60, 60, 60), after=16)
    add_kv_table(doc, [
        ("Hujjat turi", "Texnik topshiriq (TZ) va mahsulot blueprint"),
        ("Versiya", "1.0"),
        ("Status", "Ishlab chiqishga tayyorlangan boshlang'ich professional scope"),
        ("Platforma", "Hozircha faqat macOS"),
        ("Interaction modeli", "Faqat ovoz orqali: user voice yuboradi, avatar voice bilan javob beradi"),
        ("Til talabi", "O'zbek tilida ishlashi shart; avatar o'zbekcha gapirishi shart"),
        ("Avatar sifati", "Unreal Engine / MetaHuman darajasi"),
    ])
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="2E74B5", size="12", space="5")
    add_callout(
        doc,
        "Asosiy qaror",
        "Bu loyiha chat-first app emas. Matnli chat, message bubble va asosiy text input MVP scope'dan chiqariladi. Mahsulotning markazida real-time voice conversation, o'zbekcha nutq, hissiyotli avatar va tool bajaruvchi AI companion turadi.",
        fill="EAF3FF",
    )


def build_doc():
    doc = Document()
    set_document_styles(doc)
    add_cover(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_para(doc, "Loyiha maqsadi - macOS uchun professional voice-only AI companion yaratish. Foydalanuvchi assistant bilan faqat ovoz orqali muloqot qiladi. Assistant foydalanuvchi nutqini eshitadi, o'zbek tilida tushunadi, LLM orqali qaror va javob tayyorlaydi, ElevenLabs yoki Kokoro orqali ovozga aylantiradi, Unreal/MetaHuman darajasidagi avatar esa shu javobni lip-sync, yuz ifodasi va tana animatsiyasi bilan jonli tarzda qaytaradi.")
    add_para(doc, "Mahsulot Grace/Unclaw uslubidagi tajribadan o'rganilgan logicga tayanadi: local orchestration server, Unreal character runtime, WebRTC/Pixel Streaming, STT, LLM, TTS, lipsync, emotion mapping, tools va permission-controlled desktop actions. Biroq nom, dizayn, avatar, asset, prompt va implementation custom bo'ladi.")

    doc.add_page_break()
    add_heading(doc, "2. Foydalanuvchi Tomonidan Tasdiqlangan Qarorlar", 1)
    add_table(doc, ["#", "Qaror", "TZ ichidagi talqin"], [
        ("1", "Hozircha faqat macOS", "MVP macOS native desktop tajriba sifatida quriladi. Windows/Linux keyingi bosqichga qoldiriladi."),
        ("2", "Unreal/MetaHuman darajasida", "Avatar runtime boshidan high-fidelity 3D digital human sifatida rejalashtiriladi."),
        ("3", "User voice yuboradi, avatar voice bilan javob beradi", "Chat UI asosiy interaction emas. Push-to-talk/wake-listening, STT, LLM, TTS va avatar playback asosiy oqim bo'ladi."),
        ("4", "Avatar Uzbek tilida gapirishi shart", "Uzbek STT, Uzbek LLM response va Uzbek TTS majburiy acceptance criterion bo'ladi."),
        ("5", "ElevenLabs ham Kokoro ham kerak", "TTS provider abstraction boshidan ikki yo'lni qo'llab-quvvatlaydi: cloud quality va local/offline fallback."),
        ("6", "Hozircha local", "Backend, settings, memory, runtime va avatar control local-first bo'ladi. Cloud sync/payment keyin qo'shiladi."),
        ("7", "Multi-provider bo'lsin", "LLM provider layer OpenAI, Anthropic/Gemini/Groq/Ollama kabi providerlarni almashtirishga tayyor bo'ladi."),
        ("8", "O'zbek tilida ishlashi shart", "UI, system prompt, STT/TTS validation va assistant persona Uzbek-first bo'ladi."),
        ("9", "Professional bo'lsin", "UX, permission flow, logs, settings, status, packaging va recovery production-grade talabda bo'ladi."),
        ("10", "Bir nechta avatar/agent kerak", "Roster modeli: bir nechta assistant persona/avatar profillari, voice, style va permissions bilan."),
        ("11", "Internet/tools hoziroq kerak", "MVP ichida web search, browser/action tools, weather/news/stocks, reminders va permissioned desktop tool layer bo'ladi."),
    ], [600, 2600, 6160])

    add_heading(doc, "3. Scope", 1)
    add_heading(doc, "3.1 In Scope", 2)
    for item in [
        "macOS desktop app va local runtime supervisor.",
        "Unreal Engine / MetaHuman darajasidagi avatar runtime.",
        "Voice-only user interaction: push-to-talk birinchi, wake word keyingi kengaytma sifatida.",
        "Uzbek language first: STT, LLM response, TTS va persona behavior.",
        "ElevenLabs va Kokoro TTS providerlari.",
        "Multi-provider LLM routing.",
        "Local-first memory, settings, provider keys va runtime data.",
        "Tools: web search, browser action, reminders, weather, news, stocks, screen/file/computer control permission bilan.",
        "Bir nechta avatar/agent roster va persona customization.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "3.2 Out of Scope for MVP", 2)
    for item in [
        "Mobile companion app va cloud pairing.",
        "Public marketplace yoki agent store.",
        "Payment/subscription production billing.",
        "Windows/Linux support.",
        "Full custom avatar creation pipeline ichki editor bilan.",
        "Korxona darajasidagi admin panel va team management.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. Product Experience", 1)
    add_para(doc, "Foydalanuvchi appni ochganda asosiy ekran chat oynasi emas, balki tirik avatar sahnasi bo'ladi. Avatar idle holatda turadi, nafas oladi, ko'z harakatlari va kichik mikro-ifodalar bilan tiriklik hissini beradi. Foydalanuvchi mikrofon tugmasini bosib gapiradi yoki keyingi bosqichda wake word orqali assistantni chaqiradi.")
    add_table(doc, ["Holat", "Foydalanuvchi ko'radigan signal", "Ichki tizim vazifasi"], [
        ("Idle", "Avatar sokin turadi, status tayyor", "Runtime tayyor, mikrofon passiv yoki wake-listening"),
        ("Listening", "Mikrofon aktiv, avatar quloq solayotgandek", "VAD audio oqimini yig'adi"),
        ("Transcribing", "Qisqa processing status", "STT Uzbek nutqni textga aylantiradi"),
        ("Thinking", "Avatar o'ylayotgandek, ko'z/gaze harakati", "LLM, router va tool qarori ishlaydi"),
        ("Acting", "Tool bajarilayotganini bildiruvchi status", "Browser/search/reminder/system action bajariladi"),
        ("Speaking", "Avatar o'zbekcha gapiradi", "TTS, lip-sync, emotion va animation playback"),
        ("Interrupted", "Assistant to'xtab userni eshitadi", "Barge-in, playback stop, yangi utterance start"),
        ("Error", "Aniq, tinch xato holati", "Fallback yoki sozlama kerakligini bildiradi"),
    ], [1500, 3100, 4760])

    add_heading(doc, "5. High-Level Architecture", 1)
    add_callout(doc, "Arxitektura prinsipi", "Desktop shell foydalanuvchi tajribasini boshqaradi. Local orchestration server AI pipeline'ni boshqaradi. Avatar runtime 3D personajni render qiladi. Har bir qatlam mustaqil restart, status va fallbackga ega bo'lishi kerak.", fill="F4F6F9")
    add_table(doc, ["Qatlam", "Texnologiya", "Mas'uliyat"], [
        ("Desktop Shell", "Electron yoki Tauri; macOS permissions", "Window, onboarding, settings, microphone capture, local server supervision, status UI"),
        ("AI Orchestration Server", "Python FastAPI", "STT, LLM routing, TTS, lipsync, emotion mapping, memory, tools, API endpoints"),
        ("Avatar Runtime", "Unreal Engine + MetaHuman/custom digital human", "3D render, animation playback, avatar switching, lighting, camera, idle/talking states"),
        ("Streaming Bridge", "Unreal Pixel Streaming / WebRTC", "Unreal video/audio streamni desktop shell ichiga olib kelish"),
        ("Local Storage", "SQLite + encrypted key storage + file cache", "User profile, memory, reminders, provider config, logs, models cache"),
        ("Tool Layer", "MCP-compatible local tools + custom adapters", "Browser, search, screen, files, reminders, widgets, computer control"),
    ], [1800, 2500, 5060])

    add_heading(doc, "6. Voice Pipeline", 1)
    for step in [
        "Audio capture: macOS microphone permission olinadi, input stream local enginega beriladi.",
        "Pre-processing: noise suppression, echo cancellation va volume normalization qo'llanadi.",
        "VAD: user gapirishni boshlagan va tugatgan joy aniqlanadi.",
        "STT: Uzbek speech textga aylantiriladi. MVP local STT bilan boshlanadi.",
        "Utterance router: text normal suhbat, command yoki tool request ekanini aniqlaydi.",
        "LLM: provider tanlanadi, persona/memory/context qo'shiladi va structured response olinadi.",
        "TTS: response Uzbek voicega aylantiriladi. ElevenLabs va Kokoro providerlari qo'llab-quvvatlanadi.",
        "Lip-sync: audio asosida og'iz blendshape/viseme signal yaratiladi.",
        "Emotion: mood va behavior asosida yuz, gaze, blink va tana gesture parametrlari chiqariladi.",
        "Playback: avatar audio bilan sinxron gapiradi; user gapirsa barge-in orqali to'xtaydi.",
    ]:
        add_number(doc, step)

    add_heading(doc, "7. Uzbek Language Requirements", 1)
    add_table(doc, ["Yo'nalish", "Majburiy talab", "Acceptance check"], [
        ("STT", "O'zbekcha gaplarni barqaror tanishi kerak", "Oddiy suhbat, savol, buyruq va aralash Uzbek/Russian/English so'zlar test qilinadi"),
        ("LLM", "Assistant o'zbekcha tabiiy javob beradi", "Javoblar lotin o'zbek tilida, og'zaki uslubda, markdownsiz bo'ladi"),
        ("TTS", "Avatar o'zbekcha talaffuz bilan gapiradi", "ElevenLabs Uzbek voice yoki mos custom voice; Kokoro fallback sifatida baholanadi"),
        ("Persona", "Assistant userga mos o'zbekcha muomala qiladi", "Hurmat, qisqalik, tabiiylik va kontekstga mos ton tekshiriladi"),
        ("Commands", "Voice buyruqlar o'zbekcha ishlaydi", "Masalan: ob-havoni ayt, eslatma qo'y, yangiliklarni top, brauzerda och"),
    ], [1700, 3900, 3760])
    add_callout(doc, "Uzbek voice riski", "ElevenLabs Uzbek talaffuz sifati alohida sinovdan o'tkaziladi. Kokoro local fallback Uzbek uchun yetarli bo'lmasa, Kokoro faqat offline/default fallback bo'lib qoladi va MVP quality voice ElevenLabs orqali beriladi.", fill="FFF7E6")

    add_heading(doc, "8. Avatar and Unreal Requirements", 1)
    for item in [
        "Avatar sifati MetaHuman yoki unga yaqin rigged digital human darajasida bo'lishi kerak.",
        "Character roster boshidan ko'p avatarli bo'lishi kerak: har bir agentning avatar ID, voice ID, persona, style va permissions sozlamalari bo'ladi.",
        "Unreal runtime offscreen render va Pixel Streaming orqali desktop shell ichida ko'rsatiladi.",
        "Default render target 9:16 companion sahna uchun optimallashtiriladi; keyin responsive dynamic resolution qo'shiladi.",
        "Avatar idle, listening, thinking, speaking, acting va error holatlarida alohida animatsiya ko'rsatadi.",
        "Lip-sync audio bilan 60 FPS animation track sifatida sinxronlanadi, lekin render FPS 24 yoki 30 bo'lishi mumkin.",
        "Birinchi MVP bitta high-quality avatar bilan boshlanishi mumkin, lekin data model va UI bir nechta avatarni qo'llashi shart.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9. LLM Multi-Provider Layer", 1)
    add_para(doc, "Provider layer bir xil contractga ishlaydi: user utterance, profile, memory, tool context va persona kiradi; structured voice response chiqadi. Provider almashtirilganda voice pipeline va avatar pipeline o'zgarmasligi kerak.")
    add_table(doc, ["Provider turi", "MVP holati", "Izoh"], [
        ("OpenAI", "Primary candidate", "Uzbek reasoning, tool routing va structured output uchun kuchli asos"),
        ("Anthropic", "Secondary candidate", "Long context va safety kuchli, adapter orqali qo'shiladi"),
        ("Gemini", "Secondary candidate", "Web/live info va multimodal kengaytmalar uchun foydali"),
        ("Groq", "Fast/local-feeling cloud", "Tez javoblar uchun ishlatilishi mumkin"),
        ("Ollama/local", "Local-first fallback", "Offline/privacy mode, lekin Uzbek sifati alohida sinovdan o'tadi"),
    ], [1800, 2100, 5460])

    add_heading(doc, "10. Structured Response Contract", 1)
    add_table(doc, ["Maydon", "Turi", "Vazifa"], [
        ("response", "string", "Avatar ovoz bilan aytadigan yakuniy gap. Uzbek, og'zaki, markdownsiz."),
        ("mood", "string", "Yuz ifodasi va Text2Face/emotion model uchun prompt."),
        ("behavior", "enum/string", "Animation preset: listen, think, speak_happy, concerned, celebrate va hokazo."),
        ("action", "object/null", "Tool yoki command chaqiruvi. Riskli actionlarda confirmation talab qilinadi."),
        ("memory_update", "object/null", "Muhim user faktlarini local memoryga yozish taklifi."),
        ("safety_level", "enum", "normal, confirm_required, deny, sensitive."),
        ("speech_style", "enum", "brief, normal, explanatory; TTS javob uzunligini boshqaradi."),
    ], [1800, 1700, 5860])

    add_heading(doc, "11. TTS Requirements", 1)
    add_heading(doc, "11.1 ElevenLabs", 2)
    for item in [
        "Uzbek talaffuz sifati MVP acceptance uchun alohida test qilinadi.",
        "Bir nechta avatar/agent uchun alohida voice ID bo'lishi mumkin.",
        "Latency va cost monitoring kerak.",
        "API key encrypted storage orqali saqlanadi.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "11.2 Kokoro", 2)
    for item in [
        "Local/offline TTS fallback sifatida integratsiya qilinadi.",
        "Agar Uzbek quality yetarli bo'lmasa, English/Russian fallback yoki test mode sifatida belgilanadi.",
        "Provider contract ElevenLabs bilan bir xil bo'lishi kerak: text, voice_id, speed/style kiradi, PCM/WAV chiqadi.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "12. Tools and Action Layer", 1)
    add_para(doc, "Tools MVP scope'da hoziroq kerak. Biroq hammasi permission va confirmation bilan ishlaydi. Ovozli assistant foydalanuvchi nomidan real dunyo action bajarayotganda xavfsiz approval mexanizmi bo'lishi shart.")
    add_table(doc, ["Tool", "MVP talabi", "Confirmation"], [
        ("Web search", "Uzbek query bilan live internetdan javob topish", "Yo'q, read-only"),
        ("Browser control", "Sayt ochish, qidirish, sahifa o'qish", "Form submit yoki login bo'lsa kerak"),
        ("Weather", "User city/timezone bo'yicha ob-havo", "Yo'q"),
        ("News", "Headline va qisqa summary", "Yo'q"),
        ("Stocks", "Ticker narx va trend", "Yo'q, financial advice emas"),
        ("Reminders", "Voice orqali eslatma yaratish/o'zgartirish", "Create/update uchun qisqa tasdiq tavsiya etiladi"),
        ("Screen understanding", "User ruxsati bilan ekranni ko'rib izohlash", "Ha, permission kerak"),
        ("Computer control", "Click/type/open app kabi actionlar", "Ha, riskga qarab action-time confirmation"),
        ("File access", "User ruxsat bergan folderlarda qidirish/o'qish", "Sensitive file yoki upload bo'lsa ha"),
    ], [1700, 4800, 2860])

    add_heading(doc, "13. Security and Privacy", 1)
    for item in [
        "Local server faqat 127.0.0.1 interfaceda ishlaydi.",
        "API keylar macOS Keychain yoki Electron safeStorage orqali encrypted saqlanadi.",
        "Logs ichida API key, token, raw private audio va sensitive transcript saqlanmaydi.",
        "Microphone, camera, screen recording, accessibility permissionlar alohida tushuntiriladi.",
        "Computer control, file access va browser submit actionlari permission-gated bo'ladi.",
        "User 'unut' yoki 'buni saqlama' desa memory update yozilmaydi yoki mavjud memory o'chiriladi.",
        "Cloud keyingi bosqichga qoldirilgan bo'lsa ham, arxitektura keyinchalik cloud sync qo'shiladigan qilib ajratiladi.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "14. UI and UX Specification", 1)
    add_table(doc, ["Ekran/panel", "Vazifa", "Muhim elementlar"], [
        ("Main Companion", "Avatar bilan asosiy interaction", "Avatar, status pill, mic, mute, push-to-talk, active tool status"),
        ("Onboarding", "Boshlang'ich sozlash", "Permissions, Uzbek language, provider, TTS, avatar, voice test"),
        ("Provider Settings", "LLM/STT/TTS sozlash", "OpenAI/Anthropic/Gemini/Groq/Ollama, ElevenLabs/Kokoro, key status"),
        ("Avatar Roster", "Bir nechta agent/avatar", "Agent name, avatar ID, voice ID, personality, language, permissions"),
        ("Tools Permissions", "Riskli tool access boshqarish", "Browser, screen, files, computer control, confirmations"),
        ("Memory/Privacy", "Saqlangan faktlar va o'chirish", "Profile, memory, reminders, logs, clear controls"),
        ("Diagnostics", "Professional support", "Runtime health, port/status, model cache, crash logs, latency"),
    ], [1900, 2700, 4760])

    add_heading(doc, "15. Local Data Model", 1)
    add_table(doc, ["Entity", "Fields", "Izoh"], [
        ("UserProfile", "name, language, timezone, city, preferences", "Uzbek-first response va personalization uchun"),
        ("AgentProfile", "id, name, avatar_id, voice_provider, voice_id, persona, style, permissions", "Bir nechta avatar/agent roster"),
        ("ConversationSession", "id, started_at, summary, last_intents", "Chat history UI yo'q, lekin context uchun session summary bor"),
        ("MemoryItem", "id, type, content, source, confidence, created_at", "Faqat muhim va ruxsatli faktlar"),
        ("Reminder", "id, title, due_at, notes, status", "Voice command bilan boshqariladi"),
        ("ProviderConfig", "provider, model, encrypted_key_ref, status", "API keyning o'zi DBda raw saqlanmaydi"),
        ("ToolPermission", "tool_id, enabled, confirmation_policy", "Riskli actionlar uchun policy"),
        ("RuntimeHealth", "component, status, latency, last_error", "Diagnostics panel uchun"),
    ], [1800, 4300, 3260])

    doc.add_page_break()
    add_heading(doc, "16. Local API Surface", 1)
    add_table(doc, ["Endpoint", "Method", "Vazifa"], [
        ("/health", "GET", "Runtime, model, provider va avatar statusini qaytaradi"),
        ("/voice/session/start", "POST", "Voice capture yoki push-to-talk sessionni boshlaydi"),
        ("/voice/session/end", "POST", "Utterance finalization va STT trigger"),
        ("/transcribe", "POST", "Audio chunk yoki file uchun STT"),
        ("/respond", "POST", "LLM + TTS + animation pipeline"),
        ("/tts/synthesize", "POST", "Provider-based TTS"),
        ("/avatar/play", "POST", "Audio + animation jobni avatar runtimega yuborish"),
        ("/avatar/status", "GET", "Unreal/stream/player status"),
        ("/agents", "GET/POST/PATCH", "Agent roster CRUD"),
        ("/tools/execute", "POST", "Permission-gated tool action"),
        ("/settings", "GET/PATCH", "Runtime settings"),
    ], [2600, 1200, 5560])

    add_heading(doc, "17. MVP Roadmap", 1)
    add_table(doc, ["Phase", "Natija", "Exit criteria"], [
        ("Phase 0", "Blueprint, repo structure, architecture decisions", "TZ tasdiqlanadi, stack tanlanadi"),
        ("Phase 1", "macOS desktop shell + local backend supervisor", "App local serverni start/stop qiladi, health ko'rinadi"),
        ("Phase 2", "Voice capture + STT", "Uzbek utterance textga aylanadi"),
        ("Phase 3", "LLM multi-provider + Uzbek response contract", "Provider tanlanadi, structured response keladi"),
        ("Phase 4", "ElevenLabs + Kokoro TTS", "Uzbek voice playback ishlaydi, provider switch bor"),
        ("Phase 5", "Unreal/MetaHuman avatar runtime", "Avatar app ichida stream bo'ladi"),
        ("Phase 6", "Lip-sync + emotion mapping", "Avatar audio bilan sinxron gapiradi"),
        ("Phase 7", "Tools MVP", "Web search, browser, reminders, weather/news/stocks ishlaydi"),
        ("Phase 8", "Roster + settings + privacy", "Bir nechta agent sozlanadi"),
        ("Phase 9", "Polish + packaging", "macOS installable build va diagnostics tayyor"),
    ], [1500, 3900, 3960])

    add_heading(doc, "18. Acceptance Criteria", 1)
    for item in [
        "macOS app bir tugma bilan ishga tushadi va local backend statusini ko'rsatadi.",
        "User o'zbekcha voice yuboradi va assistant o'zbekcha ovoz bilan javob beradi.",
        "Text chat asosiy UI sifatida mavjud bo'lmaydi.",
        "ElevenLabs va Kokoro providerlari settings orqali tanlanadi.",
        "LLM provider layer kamida ikki provider bilan ishlaydi, interface keyin kengayadi.",
        "Unreal/MetaHuman avatar app ichida ko'rinadi va speaking holatida audio bilan sinxron animatsiya beradi.",
        "Barge-in: user assistant gapirayotgan paytda gapirsa, assistant to'xtab tinglaydi.",
        "Weather/news/stocks/reminder/web search/browser kabi tools ovozli command bilan ishlaydi.",
        "Riskli actionlarda confirmation mexanizmi bor.",
        "Logs sensitive token yoki API keylarni saqlamaydi.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "19. Risks and Mitigations", 1)
    add_table(doc, ["Risk", "Ta'sir", "Mitigation"], [
        ("Uzbek TTS sifati yetarli bo'lmasligi", "Avatar tabiiy ko'rinmaydi", "ElevenLabs custom voice sinovi, Kokoro fallback, provider switch"),
        ("Unreal Pixel Streaming murakkabligi", "MVP sekinlashadi", "Runtime'ni alohida service qilish, health/restart, minimal scene bilan boshlash"),
        ("Latency yuqori bo'lishi", "Conversation tabiiy bo'lmaydi", "Streaming STT/TTS, partial thinking animation, provider benchmark"),
        ("Tool safety", "Xavfli actionlar", "Permission policy, confirmation, audit log"),
        ("Local model disk hajmi", "Install og'irlashadi", "Lazy download, model cache, setup wizard"),
        ("Multi-provider complexity", "Scope kengayadi", "Bitta contract, provider adapterlar, staged rollout"),
    ], [2600, 2600, 4160])

    add_heading(doc, "20. Ochiq Savollar", 1)
    for item in [
        "App nomi va brand yo'nalishi qanday bo'ladi?",
        "Default avatarlar soni nechta bo'ladi: 2 ta, 3 ta yoki undan ko'p?",
        "Uzbek voice uchun erkak/ayol ovozlari alohida kerakmi?",
        "ElevenLabs custom voice yaratish uchun referens audio tayyorlanadimi?",
        "Unreal avatar assetini noldan yaratamizmi yoki MetaHuman bazasidan boshlaymizmi?",
        "Push-to-talk tugmasi keyboard shortcut bo'ladimi yoki UI tugmasi yetarlimi?",
        "Tools ichida qaysi biri birinchi demo uchun eng muhim: browser, reminders, weather yoki screen understanding?",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "21. Immediate Next Step", 1)
    add_callout(
        doc,
        "Tavsiya etilgan start",
        "Birinchi sprintda macOS desktop shell, local FastAPI backend, health/status API, provider config, push-to-talk audio capture va Uzbek STT/TTS proof-of-concept qilinadi. Unreal avatar runtime parallel ravishda minimal MetaHuman scene va Pixel Streaming proof bilan boshlanadi.",
        fill="EAF3FF",
    )

    doc.core_properties.title = "Voice-Only AI Companion TZ"
    doc.core_properties.subject = "macOS Unreal/MetaHuman Uzbek voice AI avatar technical specification"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "TZ, macOS, Unreal, MetaHuman, Uzbek, voice AI, ElevenLabs, Kokoro"
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
