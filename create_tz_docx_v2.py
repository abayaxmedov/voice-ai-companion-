from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from create_tz_docx import (
    BLUE,
    DARK_BLUE,
    INK,
    MUTED,
    LIGHT_BLUE,
    CALLOUT,
    OUT,
    add_bullet,
    add_callout,
    add_heading,
    add_kv_table,
    add_number,
    add_para,
    add_table,
    paragraph_border_bottom,
    set_cell_shading,
    set_document_styles,
    set_run_font,
    set_table_width,
)


VERSION = "2.0"


def repeat_table_header(table):
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_table_r(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = add_table(doc, headers, rows, widths, header_fill=header_fill)
    repeat_table_header(table)
    return table


def add_mono_block(doc, text, title=None):
    if title:
        add_para(doc, title, size=10, color=INK, bold=True, after=3)
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(text.strip("\n").splitlines()):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line.rstrip())
        set_run_font(run, name="Courier New", size=8.2, color=RGBColor(35, 45, 55))
    add_para(doc, "", after=5)


def add_checklist(doc, rows):
    return add_table_r(
        doc,
        ["ID", "Tekshiruv", "Qabul sharti"],
        rows,
        [900, 4200, 4260],
    )


def set_v2_furniture(doc):
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_p.add_run("Voice-Only AI Companion TZ | Implementation Blueprint")
    set_run_font(run, size=9, color=MUTED, bold=True)
    paragraph_border_bottom(header_p, color="DADCE0", size="4", space="4")

    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_p.add_run(f"Version {VERSION} | macOS MVP | Local-first")
    set_run_font(run, size=9, color=MUTED)


def add_cover_v2(doc):
    add_para(doc, "TECHNICAL SPECIFICATION / IMPLEMENTATION BLUEPRINT", size=10.5, color=MUTED, bold=True, after=8)
    title = add_para(doc, "Voice-Only AI Companion", size=25, color=RGBColor(0, 0, 0), bold=True, after=4)
    title.paragraph_format.line_spacing = 1.0
    add_para(
        doc,
        "macOS uchun Unreal/MetaHuman darajasidagi o'zbekcha ovozli AI avatar platformasi",
        size=14,
        color=RGBColor(60, 60, 60),
        after=16,
    )
    add_kv_table(
        doc,
        [
            ("Hujjat maqsadi", "Boshqa coder AI yoki developer ham shu TZ asosida 1:1 logic va behaviorni qurishi uchun to'liq blueprint"),
            ("Versiya", VERSION),
            ("Status", "Implementation-ready; yakuniy brand/nom/avatar assetlari keyin custom qilinadi"),
            ("Platforma", "Hozircha faqat macOS"),
            ("Interaction modeli", "Voice-only: user gapiradi, avatar o'zbekcha ovoz bilan javob beradi"),
            ("Vizual daraja", "Unreal Engine / MetaHuman darajasidagi high-fidelity 3D avatar"),
            ("TTS talabi", "ElevenLabs va Kokoro adapterlari; Uzbek quality gate majburiy"),
            ("Local talabi", "Runtime local-first; cloud provider faqat API key bilan explicit yoqiladi"),
            ("Muhim cheklov", "Unclaw/Grace nomi, assetlari, kodi yoki proprietary promptlari ko'chirilmaydi; faqat umumiy product logic custom quriladi"),
        ],
    )
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="2E74B5", size="12", space="5")
    add_callout(
        doc,
        "Asosiy mahsulot qarori",
        "Bu chat app emas. Asosiy ekran avatar sahnasi bo'ladi. Text chat, message bubble va doimiy text input MVPdan chiqariladi. Debug transcript faqat developer/diagnostics rejimida ko'rinishi mumkin.",
        fill="EAF3FF",
    )


def add_sources(doc):
    add_heading(doc, "1. Tekshirilgan Texnik Manbalar", 1)
    add_para(
        doc,
        "Quyidagi manbalar TZdagi texnik qarorlarni aniqlashtirish uchun ishlatildi. Ular implementation paytida qayta tekshiriladi, chunki model capability va Unreal plugin holati vaqt o'tishi bilan o'zgaradi.",
    )
    add_table_r(
        doc,
        ["Manba", "TZga ta'siri", "URL"],
        [
            (
                "Epic Games Pixel Streaming docs",
                "Pixel Streaming WebRTC orqali UE ilovasini brauzerga stream qiladi; Mac platformasi va VideoToolbox encoder qo'llab-quvvatlanishi qayd etilgan.",
                "https://dev.epicgames.com/documentation/en-us/unreal-engine/pixel-streaming-in-unreal-engine",
            ),
            (
                "Epic Pixel Streaming Reference",
                "Mac/VideoToolbox, portlar, launch arguments, encoder va WebRTC settings uchun source.",
                "https://dev.epicgames.com/documentation/unreal-engine/unreal-engine-pixel-streaming-reference",
            ),
            (
                "Epic Pixel Streaming 2 Overview",
                "UE 5.5+ uchun Pixel Streaming 2 migratsiya va mos infrastructure branch talabi.",
                "https://dev.epicgames.com/documentation/unreal-engine/pixel-streaming-2-overview-in-unreal-engine",
            ),
            (
                "ElevenLabs Models docs",
                "TTS/STT model tanlovi, Flash latency, Scribe realtime STT va supported language cheklovlari.",
                "https://elevenlabs.io/docs/overview/models",
            ),
            (
                "Kokoro-82M model card",
                "Kokoro open-weight 82M TTS, Apache-2.0, local deployment imkoniyati va voice/language cheklovlari.",
                "https://huggingface.co/hexgrad/Kokoro-82M",
            ),
        ],
        [2100, 4300, 2960],
    )
    add_callout(
        doc,
        "Uzbek TTS bo'yicha professional ogohlantirish",
        "ElevenLabs rasmiy TTS language listida Uzbek alohida ko'rsatilmagan; Kokoro voice ro'yxatida ham Uzbek yo'q. Shuning uchun TZ Uzbek voice'ni 'taxminan ishlaydi' deb emas, majburiy R&D va acceptance gate sifatida belgilaydi. MVP muvaffaqiyati ElevenLabs custom/cross-lingual voice, STT/TTS normalization va Kokoro fallback sifati bilan isbotlanishi shart.",
        fill="FFF7E6",
    )


def add_decisions(doc):
    add_heading(doc, "2. Tasdiqlangan Qarorlar", 1)
    add_table_r(
        doc,
        ["#", "Qaror", "Implementation talqini"],
        [
            ("1", "Hozircha faqat macOS", "Build, permissions, packaging va runtime supervision macOS uchun qilinadi. Windows/Linux scope emas."),
            ("2", "Unreal/MetaHuman darajasi", "Avatar pipeline boshidan high-fidelity 3D digital human sifatida quriladi; oddiy 2D avatar yoki webcam overlay yetarli emas."),
            ("3", "Voice-only", "User voice yuboradi; assistant STT -> LLM -> TTS -> avatar playback orqali javob beradi. Text chat asosiy UI emas."),
            ("4", "Avatar Uzbek gapirishi shart", "TTS, LLM response va speech normalization Uzbek-first acceptance bilan tekshiriladi."),
            ("5", "ElevenLabs ham Kokoro ham kerak", "TTSAdapter interface ikki provider bilan ishlaydi; provider switch settingsdan boshqariladi."),
            ("6", "Local", "Orchestrator, storage, settings, cache va avatar control local ishlaydi. Cloud faqat provider API sifatida explicit key bilan."),
            ("7", "Multi-provider", "LLM/STT/TTS providerlar adapter contract orqali almashtiriladi; UI va pipeline provider nomiga bog'lanmaydi."),
            ("8", "O'zbek tilida ishlashi shart", "Default prompt, UI copy, test set va acceptance Uzbek lotin yozuvida bo'ladi."),
            ("9", "Professional daraja", "Permission, logs, diagnostics, error recovery, provider status, install/update va testlar production-grade bo'ladi."),
            ("10", "Bir nechta avatar/agent", "Agent roster data modeli MVPdan boshlanadi; bitta default avatar bo'lsa ham architecture ko'p avatarni ko'taradi."),
            ("11", "Tools kerak", "Web search, browser/read, reminders, weather/news/stocks va permissioned desktop tools contractga kiradi."),
        ],
        [520, 2500, 6340],
    )


def add_product_contract(doc):
    add_heading(doc, "3. Product Contract", 1)
    add_heading(doc, "3.1 Product Promise", 2)
    add_para(
        doc,
        "Mahsulot foydalanuvchiga kompyuter ichida yashaydigan, ovoz orqali ishlaydigan, o'zbekcha gapiradigan va real desktop/tool ishlarini permission bilan bajara oladigan AI companion tajribasini beradi. Boshqa coder hujjatni olganda aynan shu tajribani ko'z oldiga keltirishi kerak: chat oynasi emas, tirik avatar bilan suhbat.",
    )
    add_heading(doc, "3.2 Non-Negotiable Requirements", 2)
    for item in [
        "App ochilganda asosiy interaction avatar va mikrofon orqali bo'ladi.",
        "Assistant javobi avatar ovozi orqali chiqadi; text response faqat diagnostics transcriptda optional.",
        "O'zbekcha tushunish va o'zbekcha gapirish release blocker hisoblanadi.",
        "Providerlar settingsdan almashtiriladi, lekin pipeline contract o'zgarmaydi.",
        "User nomidan browser, file yoki desktop action bajarishdan oldin riskga qarab confirmation olinadi.",
        "No proprietary copy: Unclaw/Grace asset, name, bundled model yoki code ko'chirilmaydi.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "3.3 MVP Out of Scope", 2)
    for item in [
        "Mobile app, public cloud pairing, paid subscription billing va team workspace.",
        "Avatar marketplace yoki in-app full character creator.",
        "Public internet hosting; MVP local desktop app sifatida ishlaydi.",
        "Text-first chatbot UI.",
        "Autonomous destructive desktop actions: delete, purchase, send, submit kabi actionlar har doim explicit confirmation talab qiladi.",
    ]:
        add_bullet(doc, item)


def add_stack(doc):
    add_heading(doc, "4. Tavsiya Etilgan Stack", 1)
    add_table_r(
        doc,
        ["Qism", "Tavsiya", "Sabab / contract"],
        [
            ("Desktop shell", "Electron + React + TypeScript", "Unclaw uslubidagi local app, WebRTC player embed, macOS permissions va Python runtime supervision uchun pragmatik."),
            ("Backend", "Python 3.11+ FastAPI + Uvicorn", "STT/LLM/TTS/lipsync tools ecosystemi kuchli; local API va job queue oson."),
            ("Contracts", "Pydantic schemas + generated TS types", "Frontend/backend drift bo'lmasligi uchun bitta schema source."),
            ("Storage", "SQLite + SQLAlchemy/Alembic", "Local-first, migration, easy backup, agent roster va memory uchun yetarli."),
            ("Secrets", "macOS Keychain yoki Electron safeStorage", "API key raw SQLite yoki logga yozilmaydi."),
            ("Audio", "Web Audio / native capture + Python audio service", "Push-to-talk, VAD, chunking, barge-in uchun."),
            ("STT", "Local Whisper/MLX first; optional ElevenLabs Scribe/OpenAI cloud", "Local-first talabga mos; cloud STT faqat user yoqsa."),
            ("LLM", "OpenAI primary adapter + Anthropic/Gemini/Groq/Ollama adapters", "Multi-provider va fallback uchun."),
            ("TTS", "ElevenLabs + Kokoro adapters", "ElevenLabs quality candidate, Kokoro local fallback/R&D."),
            ("Avatar", "Unreal Engine 5.x + MetaHuman/custom rig + Pixel Streaming 2", "High-fidelity digital human va WebRTC embed uchun."),
            ("Tools", "MCP-compatible adapters + local tool registry", "Browser/search/screen/file/computer controlni permission bilan kengaytirish uchun."),
        ],
        [1600, 3000, 4760],
    )
    add_callout(
        doc,
        "Stack lock",
        "Developer boshqa stack tanlasa ham public contractlar saqlanishi shart: local API, provider adapters, avatar bridge, state machine, storage schema va acceptance criteria o'zgarmaydi.",
        fill="F4F6F9",
    )


def add_repo_structure(doc):
    add_heading(doc, "5. Tavsiya Etilgan Repository Structure", 1)
    add_para(doc, "Bu structure coder AI uchun ishni bo'lib bajarishga mo'ljallangan. Nomlar custom bo'lishi mumkin, lekin ownership chegaralari saqlansin.")
    add_mono_block(
        doc,
        """
apps/desktop/                 macOS desktop shell, UI, WebRTC avatar view
services/orchestrator/         FastAPI, voice pipeline, providers, tools
services/avatar-bridge/        Unreal control bridge, playback jobs, health
packages/contracts/            shared DTO/schema definitions
packages/provider-tests/       STT/TTS/LLM benchmark fixtures
unreal/CompanionAvatar/        UE project, MetaHuman scene, animation BPs
assets/avatars/                avatar manifests, thumbnails, license notes
models/cache/                  local model cache, never committed
docs/                          architecture notes, QA scripts, release notes
scripts/dev/                   local start/stop/check scripts
""",
        "Repository boundary",
    )
    add_table_r(
        doc,
        ["Folder", "Owner", "Majburiy natija"],
        [
            ("apps/desktop", "Frontend/Desktop developer", "Main avatar screen, onboarding, settings, diagnostics, permission UX"),
            ("services/orchestrator", "AI/backend developer", "Voice pipeline, provider adapters, API, memory, tools, job queue"),
            ("services/avatar-bridge", "Realtime/avatar developer", "Avatar playback jobs, stream status, UE event bridge"),
            ("unreal/CompanionAvatar", "Unreal developer", "MetaHuman scene, states, lip-sync receiver, Pixel Streaming configuration"),
            ("packages/contracts", "Tech lead", "API DTOs, enums, error codes, versioned contracts"),
            ("packages/provider-tests", "QA/AI developer", "Uzbek STT/TTS/LLM regression set"),
        ],
        [2200, 2500, 4660],
    )


def add_state_machine(doc):
    add_heading(doc, "6. Runtime State Machine", 1)
    add_para(doc, "Har bir UI, backend va avatar holati bitta state machinedan keladi. Frontend alohida taxmin qilmaydi; /runtime/state yoki WebSocket eventga ishonadi.")
    add_table_r(
        doc,
        ["State", "Trigger", "UI/Avatar", "Backend action"],
        [
            ("booting", "App open", "Loading/status", "Local server, model cache, avatar runtime start"),
            ("idle", "Ready", "Avatar idle, mic ready", "Health polling, no active audio"),
            ("listening", "Push-to-talk/wake", "Mic active, avatar attentive", "Audio chunk capture + VAD"),
            ("transcribing", "Speech end", "Short processing", "STT provider request"),
            ("thinking", "Transcript ready", "Thinking animation", "LLM route, memory/context, tool decision"),
            ("confirming", "Risky action", "Voice + modal confirmation", "Wait for user yes/no"),
            ("acting", "Tool approved", "Tool status", "Execute tool, stream progress"),
            ("synthesizing", "Response ready", "Preparing speech", "TTS provider request"),
            ("speaking", "Audio ready", "Lip-sync speaking", "Avatar playback + audio output"),
            ("interrupted", "Barge-in", "Stop speaking, listen", "Cancel playback/job, start new utterance"),
            ("error", "Failure", "Calm error state", "Fallback/retry/report diagnostics"),
        ],
        [1300, 1900, 2600, 3560],
    )
    add_heading(doc, "6.1 Barge-In Contract", 2)
    for item in [
        "User assistant gapirayotganda gapirishni boshlasa, current audio playback 200 ms ichida stop bo'lishi kerak.",
        "Avatar speaking animationdan listening animationga o'tadi.",
        "Orchestrator eski turnni cancelled deb belgilaydi, lekin audit logda saqlaydi.",
        "Yangi user utterance oldingi javobni interrupt qilgani contextga kiritiladi.",
    ]:
        add_bullet(doc, item)


def add_architecture(doc):
    add_heading(doc, "7. High-Level Architecture", 1)
    add_table_r(
        doc,
        ["Layer", "Service", "Public contract"],
        [
            ("Desktop", "Electron main process", "Start/stop local backend, manage app lifecycle, macOS permissions, tray/window"),
            ("Frontend", "React renderer", "Avatar view, mic controls, settings, diagnostics, WebRTC player"),
            ("Local API", "FastAPI orchestrator", "REST + WebSocket events for voice turns, providers, agents, tools, health"),
            ("Voice pipeline", "Audio/STT/LLM/TTS services", "VoiceTurnRequest -> VoiceTurnResult contract"),
            ("Avatar bridge", "UE control bridge", "AvatarPlaybackJob -> UE event + status callbacks"),
            ("Unreal runtime", "Packaged UE app", "Pixel Streaming media + custom JSON events"),
            ("Storage", "SQLite + secure keychain", "Profiles, agents, memory, permissions, runtime logs"),
            ("Tools", "Tool registry", "Permission-scoped tool execution with audit trail"),
        ],
        [1500, 2400, 5460],
    )
    add_mono_block(
        doc,
        """
User voice
  -> Desktop audio capture
  -> Local orchestrator /voice/turn
  -> VAD + STT
  -> Intent/tool router
  -> LLM provider adapter
  -> Optional tool execution + confirmation
  -> Uzbek response contract
  -> TTS provider adapter
  -> Lip-sync/emotion mapping
  -> AvatarPlaybackJob
  -> Unreal avatar speaks through Pixel Streaming/WebRTC
""",
        "Canonical voice turn flow",
    )


def add_desktop_spec(doc):
    add_heading(doc, "8. Desktop App Specification", 1)
    add_heading(doc, "8.1 Main Companion Screen", 2)
    add_table_r(
        doc,
        ["Element", "Talab", "Behavior"],
        [
            ("Avatar viewport", "Ekranning asosiy qismi", "WebRTC/Pixel Streaming player; no chat column by default"),
            ("Mic control", "Push-to-talk/mute", "Hold/toggle modes, visual recording state, keyboard shortcut optional"),
            ("Status pill", "idle/listening/thinking/speaking/error", "Runtime state machinedan keladi"),
            ("Provider indicator", "LLM/TTS/STT health", "Small diagnostics; userga technical overload qilmaydi"),
            ("Tool status", "Action in progress", "Search/browser/reminder kabi actionlar qisqa ko'rsatiladi"),
            ("Transcript", "Hidden/dev optional", "Asosiy UI emas; diagnostics yoki accessibility mode uchun"),
            ("Settings button", "Provider/avatar/tools/privacy", "Icon button, no onboarding clutter"),
        ],
        [1800, 2800, 4760],
    )
    add_heading(doc, "8.2 Onboarding Flow", 2)
    for item in [
        "Welcome: voice-only companion konsepti, Uzbek language default.",
        "Permissions: microphone required; screen/accessibility/browser permissions optional va alohida tushuntiriladi.",
        "Provider setup: LLM, STT, TTS provider tanlash; API key encrypted storage.",
        "TTS voice test: ElevenLabs/Kokoro sample Uzbek utterance test.",
        "Avatar setup: default avatar tanlash yoki mavjud asset manifestini tanlash.",
        "Tools policy: read-only tools default on; risky tools default off.",
        "Diagnostics check: local server, model cache, avatar stream, audio input/output health.",
    ]:
        add_number(doc, item)
    add_heading(doc, "8.3 Settings Panels", 2)
    add_table_r(
        doc,
        ["Panel", "Fields", "Save behavior"],
        [
            ("Providers", "LLM provider/model/key; STT provider/model; TTS provider/voice", "Validate key, test request, encrypted save"),
            ("Voice", "Input device, output device, VAD sensitivity, barge-in", "Live test with waveform"),
            ("Avatar roster", "Agent name, avatar, voice, persona, language, style", "Save AgentProfile and VoiceProfile"),
            ("Tools", "Tool enabled, permission level, confirmation policy", "Audit permission changes"),
            ("Memory/privacy", "Saved facts, reminders, clear memory, retention", "Delete must be irreversible after confirmation"),
            ("Diagnostics", "Ports, runtime logs, model cache, stream stats", "Export safe diagnostic bundle without secrets"),
        ],
        [1700, 4300, 3360],
    )


def add_backend_spec(doc):
    add_heading(doc, "9. Local Orchestrator Backend", 1)
    add_para(doc, "Backend local FastAPI service sifatida ishlaydi va desktop shell tomonidan start/stop qilinadi. Barcha external AI providerlar adapter orqali chaqiriladi. Backend provider-specific behaviorni UIga oqizmaydi.")
    add_heading(doc, "9.1 Service Responsibilities", 2)
    for item in [
        "Audio session lifecycle va STT finalization.",
        "Conversation context, memory retrieval va prompt assembly.",
        "LLM response contract validation va retry/fallback.",
        "TTS synthesis, audio normalization va cache.",
        "Lip-sync/emotion job yaratish.",
        "Avatar bridge job dispatch va status tracking.",
        "Tool registry, permission policy, confirmation flow va audit log.",
        "Health/status/diagnostics API.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "9.2 Process Supervision", 2)
    add_table_r(
        doc,
        ["Process", "Start by", "Health check", "Restart policy"],
        [
            ("orchestrator", "Desktop main process", "/health", "Crash bo'lsa 3 marta retry, keyin diagnostics"),
            ("avatar-runtime", "Desktop or backend supervisor", "/avatar/status", "Exit/crash bo'lsa clean restart, active turn cancelled"),
            ("signalling-server", "Avatar supervisor", "/stream/status", "Port conflict aniqlansa alternative local port"),
            ("model-cache-worker", "Backend", "/models/status", "Download retry with checksum"),
            ("tool-worker", "Backend", "/tools/status", "Risky action queue persisted"),
        ],
        [1900, 2300, 2200, 2960],
    )


def add_audio_stt(doc):
    add_heading(doc, "10. Audio Capture, VAD and STT", 1)
    add_table_r(
        doc,
        ["Component", "Talab", "Acceptance"],
        [
            ("Sample rate", "16kHz or 24kHz normalized PCM internal format", "Provider adapters required formatga convert qiladi"),
            ("VAD", "Silence threshold + max utterance duration", "Short pause bilan turn tugamaydi; long silence turnni final qiladi"),
            ("Noise handling", "Basic normalization/noise suppression", "Laptop fan/background noise testdan o'tadi"),
            ("Barge-in", "Playback paytida input energy/VAD trigger", "200 ms ichida speaking cancel"),
            ("STT local", "Whisper/MLX or equivalent local Uzbek-capable model", "Uzbek test set WER threshold bilan baholanadi"),
            ("STT cloud optional", "ElevenLabs Scribe/OpenAI/other", "User enabled bo'lsa; key encrypted"),
            ("Transcript policy", "Raw transcript local va retention controlled", "Sensitive transcript default uzoq saqlanmaydi"),
        ],
        [1800, 4200, 3360],
    )
    add_heading(doc, "10.1 Uzbek STT Test Utterances", 2)
    for item in [
        "Salom, bugun Toshkentda ob-havo qanday?",
        "Ertaga soat to'qqizga menga uchrashuv eslatmasini qo'y.",
        "Chrome'da Fotonlabs saytini ochib, Grace nima ekanini tushuntir.",
        "Mening ismim Aziz, lekin buni xotiraga saqlama.",
        "Ovozingni biroz sekinroq va xotirjamroq qil.",
    ]:
        add_bullet(doc, item)


def add_llm_provider(doc):
    add_heading(doc, "11. LLM Multi-Provider Layer", 1)
    add_table_r(
        doc,
        ["Provider", "Role", "Required adapter behavior"],
        [
            ("OpenAI", "Primary cloud candidate", "Structured JSON response, tool routing, Uzbek system prompt"),
            ("Anthropic", "Secondary cloud candidate", "Same response schema; tool calls normalized"),
            ("Gemini", "Secondary/cloud multimodal candidate", "Same schema; future screen/image context"),
            ("Groq", "Low-latency candidate", "Fast text response; tool routing validated"),
            ("Ollama/local", "Local fallback", "Offline/privacy mode; Uzbek quality measured, not assumed"),
        ],
        [1600, 2500, 5260],
    )
    add_heading(doc, "11.1 Provider Adapter Contract", 2)
    add_table_r(
        doc,
        ["Input", "Meaning"],
        [
            ("utterance_text", "Final STT transcript"),
            ("language", "Default uz-Latn; mixed Uzbek/Russian/English accepted"),
            ("agent_profile", "Persona, style, allowed tools, voice constraints"),
            ("user_profile", "Name, locale, city, preferences, consented facts"),
            ("memory_context", "Only relevant, allowed memory items"),
            ("tool_context", "Available tools and permission policy"),
            ("conversation_state", "Current state, interrupted flag, previous turn summary"),
        ],
        [2500, 6860],
    )
    add_heading(doc, "11.2 Output Contract", 2)
    add_mono_block(
        doc,
        """
response: Uzbek text that avatar will speak. No markdown.
mood: neutral | happy | thoughtful | concerned | excited | apologetic
behavior: idle | listen | think | speak | explain | celebrate | confirm
speech_style: brief | normal | explanatory
action: optional tool request with tool_id, params, risk_level
confirmation_prompt: required when risk_level is medium/high
memory_update: optional fact with user consent status
safety_level: normal | confirm_required | deny | sensitive
debug_reason: developer-only short reason, never spoken
""",
        "LLMResponse shape",
    )
    add_callout(
        doc,
        "LLM retry rule",
        "Agar provider schema buzsa, backend bir marta repair prompt bilan retry qiladi. Ikkinchi xatoda fallback providerga o'tadi yoki userga ovoz bilan qisqa xato aytiladi.",
        fill=CALLOUT,
    )


def add_tts(doc):
    add_heading(doc, "12. TTS Provider Contract", 1)
    add_para(doc, "TTS layer avatar ovozining markazi. Har bir provider bir xil contractga ishlashi shart: normalized Uzbek text kiradi, audio file/stream, timing metadata va optional emotion support chiqadi.")
    add_table_r(
        doc,
        ["Provider", "MVP role", "Uzbek reality check", "Fallback"],
        [
            ("ElevenLabs", "High-quality cloud voice candidate", "Rasmiy language list Uzbekni alohida ko'rsatmaydi; custom/cross-lingual voice sinovi shart", "Agar talaffuz past bo'lsa, Uzbek-specific custom voice yoki alternative provider R&D"),
            ("Kokoro", "Local/offline/open-weight fallback", "Kokoro-82M voice listida Uzbek yo'q; G2P/phoneme limitation kutiladi", "Dev/offline mode, English/Russian/Turkic experiment, yoki future fine-tune"),
        ],
        [1600, 2500, 3260, 2000],
    )
    add_heading(doc, "12.1 Text Normalization", 2)
    for item in [
        "Raqamlar TTSga berilishdan oldin o'zbekcha yozuvga normalizatsiya qilinadi: 9:30 -> to'qqiz yarim.",
        "URL, ticker, email va qisqartmalar speakable formga o'tkaziladi.",
        "Assistant javobi uzun bo'lsa, 1-2 gaplik audio chunklarga bo'linadi.",
        "Markdown, bullets, code, table text TTSga yuborilmaydi; speaking response plain Uzbek bo'ladi.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "12.2 Voice Profile Model", 2)
    add_table_r(
        doc,
        ["Field", "Example", "Purpose"],
        [
            ("voice_profile_id", "aziza_default", "Agent voice identity"),
            ("provider", "elevenlabs | kokoro", "Provider selection"),
            ("voice_id", "provider-specific", "External/local voice reference"),
            ("language_mode", "uz-Latn", "Default speaking language"),
            ("speed", "0.85-1.10", "Natural Uzbek pace"),
            ("emotion_support", "true/false", "Can provider handle style/emotion prompts"),
            ("fallback_chain", "elevenlabs -> kokoro -> silent_error", "Recovery path"),
        ],
        [2100, 2500, 4760],
    )
    add_heading(doc, "12.3 TTS Acceptance", 2)
    add_checklist(
        doc,
        [
            ("TTS-1", "O'zbekcha 20 ta test gap tabiiy eshitiladi", "Human review: kamida 4/5 MOS-like baho"),
            ("TTS-2", "Numbers/dates normalizatsiya to'g'ri", "9:30, 12000 so'm, 2026 yil cases"),
            ("TTS-3", "Latency conversational", "Short response audio start target < 1200 ms cloud, < 700 ms local if possible"),
            ("TTS-4", "Provider switch runtime settingsdan ishlaydi", "Restart qilmasdan voice test qayta ishlaydi"),
            ("TTS-5", "Raw API key/audio secrets logga tushmaydi", "Log scan pass"),
        ],
    )


def add_avatar(doc):
    add_heading(doc, "13. Unreal / MetaHuman Avatar Runtime", 1)
    add_para(doc, "Avatar runtime mahsulotning ko'rinadigan yuragi. MVP minimal sahna bilan boshlanishi mumkin, lekin render va rig darajasi oddiy 2D avatar emas, MetaHuman/custom digital human darajasida bo'lishi kerak.")
    add_table_r(
        doc,
        ["Qism", "Talab", "Implementation note"],
        [
            ("UE version", "UE 5.5+ tavsiya; Pixel Streaming 2 branch mosligi tekshiriladi", "Exact version repo READMEda lock qilinadi"),
            ("Character", "MetaHuman yoki MetaHuman-level rig", "Face blendshapes, jaw, eyes, brows, head/neck, body idle"),
            ("Scene", "Companion framing, soft lighting, neutral background", "Desktop app ichida avatar o'qilishi kerak"),
            ("Stream", "Pixel Streaming/WebRTC local", "Mac support VideoToolbox bilan tekshiriladi"),
            ("Bridge", "JSON events to UE", "play_audio, set_state, set_emotion, interrupt, switch_avatar"),
            ("Audio", "Avatar speaking audio desktopga chiqadi", "A/V sync measured"),
            ("Fallback", "If stream fails, diagnostics + retry", "Blank viewport release blocker"),
        ],
        [1700, 3800, 3860],
    )
    add_heading(doc, "13.1 Launch Contract", 2)
    add_mono_block(
        doc,
        """
Required UE launch intent:
- connect to local signalling server on 127.0.0.1
- use Pixel Streaming / Pixel Streaming 2 plugin
- offscreen render preferred for desktop embed stability
- fixed initial resolution, then adaptive settings
- stdout logs routed to diagnostics
""",
        "Unreal runtime launch requirements",
    )
    add_heading(doc, "13.2 Avatar Event Contract", 2)
    add_table_r(
        doc,
        ["Event", "Direction", "Payload"],
        [
            ("avatar.ready", "UE -> backend", "avatar_id, capabilities, stream_id"),
            ("avatar.state", "backend -> UE", "idle/listening/thinking/speaking/acting/error"),
            ("avatar.play", "backend -> UE", "audio_url/local_path, visemes, emotion, behavior"),
            ("avatar.interrupt", "backend -> UE", "turn_id, reason"),
            ("avatar.completed", "UE -> backend", "turn_id, playback_ms, dropped_frames"),
            ("avatar.error", "UE -> backend", "code, message, recoverable"),
            ("avatar.switch", "backend -> UE", "avatar_id, agent_profile_id"),
        ],
        [1900, 2000, 5460],
    )
    add_heading(doc, "13.3 Lip-Sync and Emotion", 2)
    for item in [
        "MVP lipsync audio amplitude + phoneme/viseme estimation bilan boshlanishi mumkin, lekin jaw/lips must be synchronized.",
        "Target representation: time_ms, viseme_name, weight array.",
        "Emotion mapping LLM mood + response styledan keladi.",
        "Thinking/listening states silent animation bilan alohida bo'ladi.",
        "Speaking paytida blink, gaze va head motion natural randomization bilan ishlaydi.",
    ]:
        add_bullet(doc, item)


def add_agents_memory(doc):
    add_heading(doc, "14. Multi-Agent / Avatar Roster", 1)
    add_para(doc, "Bir nechta avatar/agent talabi MVP data modelga kiradi. Birinchi release bitta default avatar bilan boshlansa ham, UI va backend ko'p agentni ko'tarishi shart.")
    add_table_r(
        doc,
        ["Entity", "Majburiy fieldlar", "Behavior"],
        [
            ("AgentProfile", "id, name, persona, language, avatar_id, voice_profile_id, enabled_tools", "Active agent conversation style va permissionsni belgilaydi"),
            ("AvatarManifest", "avatar_id, ue_asset_ref, thumbnail, rig_type, capabilities", "UE runtime qaysi avatarni yuklashini bildiradi"),
            ("VoiceProfile", "provider, voice_id, speed, language_mode, fallback_chain", "Agent ovozini boshqaradi"),
            ("PersonaPreset", "tone, response_length, humor_level, boundaries", "LLM system promptga compile qilinadi"),
        ],
        [1800, 4200, 3360],
    )
    add_heading(doc, "14.1 Memory Rules", 2)
    for item in [
        "Memory default minimal bo'ladi: faqat foydali, ruxsatli va long-term faktlar saqlanadi.",
        "User 'buni saqlama' desa memory update yozilmaydi.",
        "User 'meni unut' desa user profile va memory clear flow ishga tushadi.",
        "Raw audio long-term saqlanmaydi; transcript retention user settingsga bog'liq.",
        "Memory har bir agentga umumiy user facts va agent-specific preferences sifatida ajratilishi mumkin.",
    ]:
        add_bullet(doc, item)


def add_data_model(doc):
    add_heading(doc, "15. Local Data Model", 1)
    add_table_r(
        doc,
        ["Table", "Core fields", "Notes"],
        [
            ("user_profiles", "id, display_name, locale, timezone, city, created_at, updated_at", "Single-user MVP, multi-user future-ready"),
            ("agent_profiles", "id, name, persona, avatar_id, voice_profile_id, active, tool_policy_id", "Roster support"),
            ("voice_profiles", "id, provider, voice_id, language_mode, speed, fallback_chain, test_status", "ElevenLabs/Kokoro switch"),
            ("provider_configs", "id, provider, type, model, encrypted_key_ref, enabled, last_validated_at", "No raw API keys"),
            ("conversation_sessions", "id, agent_id, started_at, ended_at, summary, last_state", "No chat UI required"),
            ("conversation_turns", "id, session_id, transcript, response, state, cancelled, latency_json", "Diagnostics and context"),
            ("memory_items", "id, user_id, agent_id, type, content, confidence, consent, source_turn_id", "Permissioned memory"),
            ("tool_permissions", "id, tool_id, enabled, risk_level, confirmation_policy", "Safety policy"),
            ("tool_runs", "id, tool_id, turn_id, params_redacted, result_summary, status", "Audit log"),
            ("reminders", "id, title, due_at, timezone, status, created_by_turn_id", "Voice-created reminders"),
            ("avatar_assets", "id, name, manifest_path, ue_asset_ref, license_status", "No proprietary copied assets"),
            ("runtime_events", "id, component, level, code, message_redacted, created_at", "Diagnostics without secrets"),
        ],
        [1900, 4600, 2860],
    )
    add_callout(
        doc,
        "Migration rule",
        "Har bir DB schema change migration bilan bo'ladi. Developer DBni qo'lda buzmasligi kerak; app startup migration statusni tekshiradi.",
        fill=CALLOUT,
    )


def add_api_surface(doc):
    add_heading(doc, "16. Local API Surface", 1)
    add_para(doc, "All endpoints local-only loopback serverda ishlaydi. Frontend API token/session bilan ulanadi; CORS faqat app origin/localhostga cheklanadi.")
    add_table_r(
        doc,
        ["Endpoint", "Method", "Purpose", "Acceptance"],
        [
            ("/health", "GET", "Backend, providers, DB, model cache status", "App startup status shows ready/degraded/error"),
            ("/runtime/state", "GET/WS", "Canonical state machine events", "UI never invents state locally"),
            ("/voice/session/start", "POST", "Begin push-to-talk/listening", "Returns session_id and audio config"),
            ("/voice/session/chunk", "POST/WS", "Send audio chunks", "Handles streaming and backpressure"),
            ("/voice/session/end", "POST", "Finalize utterance", "Triggers STT and voice turn"),
            ("/voice/turn", "POST", "Full voice turn orchestration", "Returns turn_id; progress via WS"),
            ("/stt/transcribe", "POST", "File/chunk transcription", "Provider-independent transcript result"),
            ("/llm/respond", "POST", "LLM response contract", "Validates schema and fallback"),
            ("/tts/synthesize", "POST", "Provider TTS audio generation", "Returns audio ref and timing metadata"),
            ("/avatar/play", "POST", "Dispatch AvatarPlaybackJob", "UE starts speaking or returns error"),
            ("/avatar/status", "GET", "Stream/runtime/avatar health", "Used by diagnostics and auto-restart"),
            ("/agents", "GET/POST/PATCH", "Agent roster CRUD", "Validates avatar/voice references"),
            ("/tools", "GET", "Available tool registry", "Includes permission and risk"),
            ("/tools/execute", "POST", "Run permissioned tool", "Blocks risky action without confirmation"),
            ("/memory", "GET/DELETE", "Memory review and deletion", "User can inspect/clear saved facts"),
            ("/settings", "GET/PATCH", "Runtime/provider settings", "Save validates and audits"),
            ("/diagnostics/export", "POST", "Safe diagnostic bundle", "Secrets/audio stripped"),
        ],
        [2100, 1100, 3100, 3060],
    )
    add_heading(doc, "16.1 Important DTOs", 2)
    add_mono_block(
        doc,
        """
VoiceTurnRequest:
  session_id, agent_id, audio_ref or transcript_override
  interrupt_previous: boolean
  user_locale: uz-Latn

VoiceTurnResult:
  turn_id, transcript, spoken_response
  action_status, audio_ref, avatar_job_id
  latency: stt_ms, llm_ms, tts_ms, avatar_ms

AvatarPlaybackJob:
  job_id, turn_id, avatar_id
  audio_ref, visemes, mood, behavior
  allow_interrupt: true
""",
        "DTO summary",
    )


def add_tools(doc):
    add_heading(doc, "17. Tools and Permission Layer", 1)
    add_table_r(
        doc,
        ["Tool", "MVP capability", "Risk", "Confirmation policy"],
        [
            ("web_search", "Live internet search and summary in Uzbek", "Low", "No confirmation for read-only"),
            ("browser_read", "Open/read page, summarize visible content", "Low", "No confirmation unless login/private content"),
            ("browser_action", "Click/type/navigation", "Medium", "Confirm before submit/login/purchase/send"),
            ("weather", "Location-based weather", "Low", "No confirmation"),
            ("news", "Recent headlines + short Uzbek summary", "Low", "No confirmation"),
            ("stocks", "Ticker price/trend only", "Medium", "No investment advice; disclaim if needed"),
            ("reminders", "Create/update/delete reminders", "Medium", "Confirm create/update/delete"),
            ("screen_read", "User-approved screenshot/screen explanation", "Medium", "Ask permission each session"),
            ("file_read", "Read allowed folders/files", "Medium", "Confirm sensitive files"),
            ("computer_control", "Click/type/open app", "High", "Action-time confirmation"),
        ],
        [1700, 3300, 1200, 3160],
    )
    add_heading(doc, "17.1 Tool Execution Rules", 2)
    for item in [
        "Tool call LLMdan kelsa ham backend permission policy tekshirmasdan bajarmaydi.",
        "High-risk tool calls voice confirmation talab qiladi: assistant o'zbekcha nima qilmoqchi ekanini aytadi va user 'ha' demaguncha action bajarilmaydi.",
        "Tool result LLMga qisqa, redacted summary sifatida qaytariladi.",
        "Tool audit logda tool_id, risk, redacted params, status va turn_id saqlanadi.",
        "Browser login, payment, message send, file delete kabi actionlar MVPda default disabled yoki confirm_each bo'ladi.",
    ]:
        add_bullet(doc, item)


def add_security(doc):
    add_heading(doc, "18. Security, Privacy and Compliance", 1)
    add_table_r(
        doc,
        ["Area", "Requirement", "Implementation detail"],
        [
            ("Network", "Local-only by default", "Server binds 127.0.0.1; no public listener"),
            ("Auth", "Local API session token", "Desktop receives ephemeral token at backend startup"),
            ("Secrets", "No raw API key in DB/logs", "Keychain/safeStorage encrypted refs only"),
            ("Audio", "Raw audio retention off by default", "Temp files deleted after turn unless debug mode"),
            ("Transcript", "Retention controlled", "User can clear transcript/memory"),
            ("Permissions", "macOS permission explanations", "Microphone required; screen/accessibility optional"),
            ("Logs", "Redacted diagnostics", "Tokens, paths, private transcript snippets redacted"),
            ("Model cache", "Checksums and license notes", "Downloaded local models tracked with manifest"),
            ("Actions", "Risk-gated tools", "Confirmation and audit trail"),
        ],
        [1700, 3100, 4560],
    )
    add_checklist(
        doc,
        [
            ("SEC-1", "Local API external networkdan ko'rinmaydi", "Port scan from LAN fails"),
            ("SEC-2", "API key logga tushmaydi", "Automated log scan pass"),
            ("SEC-3", "Permission prompts understandable Uzbek/English", "User can deny optional tools without app crash"),
            ("SEC-4", "Memory clear works", "DB rows removed and app state updated"),
            ("SEC-5", "Risky tool action confirmation works", "No browser submit without explicit yes"),
        ],
    )


def add_diagnostics(doc):
    add_heading(doc, "19. Diagnostics and Error Handling", 1)
    add_table_r(
        doc,
        ["Error code", "User-facing Uzbek message", "Developer action"],
        [
            ("AUDIO_NO_MIC", "Mikrofon ruxsati kerak.", "Open macOS permission guidance"),
            ("STT_FAILED", "Gapingizni tushunmadim, yana bir bor ayting.", "Retry STT, log provider error redacted"),
            ("LLM_SCHEMA_INVALID", "Javob tayyorlashda xatolik bo'ldi.", "Repair retry then fallback provider"),
            ("TTS_UNAVAILABLE", "Ovoz yaratishda muammo bo'ldi.", "Switch fallback provider or ask settings"),
            ("AVATAR_STREAM_DOWN", "Avatar stream hozir ulanmagan.", "Restart UE/signalling, show diagnostics"),
            ("TOOL_CONFIRM_REQUIRED", "Buni qilishim uchun tasdiqlang.", "Wait for yes/no turn"),
            ("PROVIDER_KEY_INVALID", "Provider API key noto'g'ri yoki ishlamayapti.", "Open provider settings test"),
            ("MODEL_CACHE_MISSING", "Local model hali tayyor emas.", "Download/checksum workflow"),
        ],
        [1900, 3500, 3960],
    )
    add_heading(doc, "19.1 Health Dashboard", 2)
    for item in [
        "Backend status: ready/degraded/error.",
        "Provider status: STT, LLM, TTS test result and last latency.",
        "Avatar status: UE process, signalling server, WebRTC player, current avatar.",
        "Audio status: input/output device, VAD, last capture level.",
        "Model cache: installed/missing/downloading/checksum failed.",
        "Tool status: enabled, permission state, last run.",
    ]:
        add_bullet(doc, item)


def add_packaging(doc):
    add_heading(doc, "20. macOS Packaging and Local Install", 1)
    add_table_r(
        doc,
        ["Package item", "Requirement", "Notes"],
        [
            ("App bundle", ".app packaged desktop shell", "Signed/notarized later; dev build first"),
            ("Python runtime", "Bundled or managed virtualenv", "Version locked, dependency install progress shown"),
            ("Unreal runtime", "Packaged UE app inside resources or downloaded asset", "Large asset handling via manifest/cache"),
            ("Model downloads", "Lazy download with checksum", "User sees progress and disk estimate"),
            ("Config location", "Application Support/<AppName>", "DB, settings, logs, cache separated"),
            ("Uninstall", "Remove app; optional clear data button", "Never silently delete user data"),
            ("Updates", "Manual/dev first", "Auto-update future scope"),
        ],
        [1800, 3300, 4260],
    )
    add_callout(
        doc,
        "macOS hardware note",
        "Pixel Streaming on Mac depends on VideoToolbox-capable hardware and target macOS testing. MVP must validate Apple Silicon Mac first. If Pixel Streaming fails on a target Mac, release is blocked until fallback or requirement change is approved.",
        fill="FFF7E6",
    )


def add_qa(doc):
    add_heading(doc, "21. Acceptance Criteria and Test Plan", 1)
    add_heading(doc, "21.1 End-to-End Acceptance", 2)
    add_checklist(
        doc,
        [
            ("E2E-1", "App opens and avatar appears", "No blank viewport; /health ready"),
            ("E2E-2", "User speaks Uzbek", "Transcript is correct enough for command"),
            ("E2E-3", "Assistant responds Uzbek by voice", "Avatar audio output Uzbek and natural enough"),
            ("E2E-4", "Avatar lip-syncs response", "Mouth movement aligned with audio, no frozen face"),
            ("E2E-5", "Barge-in works", "Assistant stops speaking and listens"),
            ("E2E-6", "Provider switch works", "ElevenLabs/Kokoro can be selected and tested"),
            ("E2E-7", "Multi-provider LLM works", "At least two LLM adapters pass same response contract"),
            ("E2E-8", "Tools work by voice", "Search/weather/reminder/browser read demos pass"),
            ("E2E-9", "Risk confirmation works", "Browser submit/computer control blocked until yes"),
            ("E2E-10", "No chat-first UI", "Main screen remains avatar/voice focused"),
        ],
    )
    add_heading(doc, "21.2 Uzbek Regression Script", 2)
    add_table_r(
        doc,
        ["Scenario", "User says", "Expected assistant behavior"],
        [
            ("Weather", "Bugun Toshkentda ob-havo qanday?", "Uzbek voice summary with temperature and short advice"),
            ("Reminder", "Ertaga soat 9 ga eslatma qo'y: shifokorga borish.", "Asks confirmation or creates reminder and repeats details"),
            ("Browser read", "Fotonlabs saytini ochib, Grace haqida qisqa tushuntir.", "Opens/reads page and summarizes in Uzbek"),
            ("Memory deny", "Mening ismim Aziz, lekin buni saqlama.", "Replies naturally and does not write memory"),
            ("Interrupt", "Assistant speaking while user says: to'xta", "Stops audio and asks/listens"),
            ("Provider failure", "TTS provider disabled", "Falls back or speaks error without crash"),
            ("Risk action", "Chrome'da login formni to'ldir.", "Explains confirmation needed; does not submit automatically"),
        ],
        [1800, 3300, 4260],
    )
    add_heading(doc, "21.3 Performance Targets", 2)
    add_table_r(
        doc,
        ["Metric", "MVP target", "Stretch target"],
        [
            ("App cold start to ready", "< 30 sec after dependencies installed", "< 10 sec"),
            ("STT final after speech end", "< 1500 ms", "< 700 ms"),
            ("LLM first response", "< 2500 ms cloud", "< 1200 ms"),
            ("TTS audio start", "< 1200 ms cloud", "< 700 ms"),
            ("Barge-in stop", "< 200 ms", "< 100 ms"),
            ("Avatar stream first frame", "< 10 sec runtime warm", "< 3 sec"),
            ("A/V sync", "Visible acceptable", "< 120 ms offset"),
        ],
        [2600, 3300, 3460],
    )


def add_roadmap(doc):
    add_heading(doc, "22. Implementation Roadmap", 1)
    add_table_r(
        doc,
        ["Phase", "Build", "Exit criteria"],
        [
            ("0", "Repo scaffold, contracts, dev scripts, architecture lock", "All folders and DTO names exist; no UI yet"),
            ("1", "Desktop shell + local backend health", "App starts backend and shows health"),
            ("2", "Audio capture + VAD + STT", "Uzbek utterance transcript appears in diagnostics"),
            ("3", "LLM provider adapter + response contract", "Two providers pass same schema tests"),
            ("4", "ElevenLabs + Kokoro TTS adapters", "Voice test works and provider switch persists"),
            ("5", "Unreal/MetaHuman runtime + stream", "Avatar visible in app via WebRTC/Pixel Streaming"),
            ("6", "Avatar playback + lip-sync + emotion", "Voice response plays through avatar with state changes"),
            ("7", "Full voice turn orchestration", "User voice -> avatar Uzbek response end-to-end"),
            ("8", "Tools MVP + confirmations", "Search/weather/reminder/browser demos pass"),
            ("9", "Agent roster + settings + privacy", "Multiple agents/profiles can be configured"),
            ("10", "Diagnostics, packaging, QA hardening", "Installable macOS build and test checklist pass"),
        ],
        [900, 4300, 4160],
    )
    add_heading(doc, "22.1 Recommended First Sprint", 2)
    add_callout(
        doc,
        "Start here",
        "Birinchi sprintda desktop shell, local FastAPI health, provider settings, microphone capture, Uzbek STT proof, ElevenLabs/Kokoro TTS smoke test va minimal avatar runtime health check qilinsin. Unreal stream parallel branchda yuradi, lekin API contract oldindan yoziladi.",
        fill="EAF3FF",
    )


def add_coder_handoff(doc):
    add_heading(doc, "23. Coder AI Handoff Rules", 1)
    add_para(doc, "Agar bu hujjat boshqa coder AIga berilsa, quyidagi qoidalar promptga qo'shiladi. Bu bo'limning maqsadi implementation driftni oldini olish.")
    for item in [
        "Avval repo structure va shared contractsni yarat, keyin feature implement qil.",
        "Chat-first UI qurma. Main screen avatar + voice bo'lsin.",
        "Provider logicni UIga hardcode qilma; adapter interface ishlat.",
        "O'zbek TTS sifatini taxmin qilma; test set va human review gate qo'y.",
        "Unreal/MetaHuman assetlarini custom yarat yoki legal/licensed asset ishlat; Unclaw/Grace assetlarini ko'chirma.",
        "Har bir risky tool action uchun confirmation policy implement qil.",
        "Logs va diagnosticsdan API key/raw private audio/sensitive transcriptni chiqarib tashla.",
        "Har bir phase uchun automated yoki manual acceptance check yoz.",
        "Agar Pixel Streaming macOS targetda ishlamasa, fallbackni hujjatlashtir va release blocker sifatida belgilashdan oldin user approval so'ra.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "23.1 Definition of Done", 2)
    add_checklist(
        doc,
        [
            ("DOD-1", "End-to-end voice turn works", "User Uzbek voice -> avatar Uzbek voice response"),
            ("DOD-2", "Avatar is high fidelity", "Unreal/MetaHuman level, no placeholder cartoon in final MVP"),
            ("DOD-3", "No chat dependency", "Core use works without typing"),
            ("DOD-4", "Providers are replaceable", "TTS/LLM/STT adapters tested"),
            ("DOD-5", "Security gates pass", "Secrets, permissions, risky actions handled"),
            ("DOD-6", "Diagnostics usable", "Developer can see why STT/TTS/avatar failed"),
            ("DOD-7", "Uzbek QA passes", "Regression script accepted by user/human reviewer"),
        ],
    )


def add_risks_open_questions(doc):
    add_heading(doc, "24. Risks, Decisions and Open Questions", 1)
    add_table_r(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ("Uzbek TTS quality", "Main product promise fails", "Custom ElevenLabs voice, alternative provider research, Kokoro limited fallback, human MOS gate"),
            ("Pixel Streaming on macOS instability", "Avatar cannot embed reliably", "Validate early on target Mac, lock UE version, consider alternate local render path"),
            ("Latency too high", "Conversation feels unnatural", "Streaming STT/TTS, shorter responses, fast provider, thinking animation"),
            ("Scope creep from tools", "MVP delayed", "Read-only tools first, high-risk tools staged"),
            ("Large Unreal assets", "Install heavy", "Lazy download, asset manifest, disk estimate"),
            ("Provider API cost", "Usage expensive", "Local fallback, usage limits, per-provider metrics"),
            ("Safety/privacy", "Trust loss", "Permission gates, redacted logs, memory controls"),
        ],
        [2400, 3000, 3960],
    )
    add_heading(doc, "24.1 Open Questions", 2)
    for item in [
        "Appning final nomi va brand tone qanday bo'ladi?",
        "Default avatarlar soni: 1, 2 yoki 3 ta?",
        "Uzbek voice uchun erkak/ayol variantlari kerakmi?",
        "ElevenLabs custom voice uchun legal referens audio tayyorlanadimi?",
        "Target Mac modeli: Apple Silicon M1/M2/M3/M4 qaysi minimum?",
        "Wake word MVPga kiradimi yoki push-to-talk yetarlimi?",
        "Tools demo priority: web search, browser read, reminders, weather, screen understanding qaysi birinchi?",
    ]:
        add_bullet(doc, item)


def build_doc():
    doc = Document()
    set_document_styles(doc)
    set_v2_furniture(doc)
    add_cover_v2(doc)

    add_sources(doc)
    doc.add_page_break()
    add_decisions(doc)
    add_product_contract(doc)
    add_stack(doc)
    doc.add_page_break()
    add_repo_structure(doc)
    add_state_machine(doc)
    add_architecture(doc)
    add_desktop_spec(doc)
    doc.add_page_break()
    add_backend_spec(doc)
    add_audio_stt(doc)
    doc.add_page_break()
    add_llm_provider(doc)
    add_tts(doc)
    add_avatar(doc)
    add_agents_memory(doc)
    doc.add_page_break()
    add_data_model(doc)
    doc.add_page_break()
    add_api_surface(doc)
    doc.add_page_break()
    add_tools(doc)
    add_security(doc)
    add_diagnostics(doc)
    add_packaging(doc)
    add_qa(doc)
    add_roadmap(doc)
    add_coder_handoff(doc)
    add_risks_open_questions(doc)

    doc.core_properties.title = "Voice-Only AI Companion TZ Implementation Blueprint"
    doc.core_properties.subject = "macOS Unreal MetaHuman Uzbek voice AI avatar implementation-ready technical specification"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "TZ, implementation blueprint, macOS, Unreal, MetaHuman, Uzbek, voice AI, ElevenLabs, Kokoro"
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
